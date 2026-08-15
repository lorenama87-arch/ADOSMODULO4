import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- CONFIGURACIÓN DE PÁGINA (DEBE IR AQUÍ, ANTES DE CUALQUIER OTRA COSA) ---
st.set_page_config(page_title="ADOS-2 Master M4", layout="wide")

# --- SISTEMA DE SEGURIDAD DE LA APP ---
def check_password():
    """Devuelve True si el usuario ingresó la contraseña correcta."""
    def password_entered():
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("🔐 Introduce la contraseña de acceso:", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("🔐 Introduce la contraseña de acceso:", type="password", on_change=password_entered, key="password")
        st.error("Contraseña incorrecta. Acceso denegado.")
        return False
    else:
        return True

if not check_password():
    st.stop()

# 1. LÓGICA DE PUNTUACIÓN (sin cambios: es el algoritmo oficial, no se toca)
def c(v):
    """Conversión oficial del algoritmo (3->2, 7/8/9->0)"""
    if v in [0, 7, 8, 9]: return 0
    if v == 1: return 1
    if v in [2, 3]: return 2
    return 0

# 2. CONSTANTES Y DICCIONARIOS
TASKS = ["Tarea de construcción", "Contar una historia de un libro", "Descripción de una imagen", "Conversación y narración", "Trabajo o escuela", "Dificultades sociales", "Emociones", "Tarea de demostración", "Viñetas", "Descanso", "Vida diaria", "Amistades y matrimonio", "Soledad", "Planes e ilusiones", "Inventarse una historia"]
ALGO_KEYS = {"A4", "A8", "A9", "A10", "B1", "B2", "B6", "B8", "B9", "B11", "B12"}
UI_STRUCTURE = {
    "A. Lenguaje": [("A1", [0,1,2,3]), ("A2", [0,1,2,7]), ("A3", [0,1,2,3]), ("A4", [0,1,2,3]), ("A5", [0,1,2]), ("A6", [0,1,2,3]), ("A7", [0,1,2,3]), ("A8", [0,1,2,3]), ("A9", [0,1,2,3,8]), ("A10", [0,1,2,3,8])],
    "B. Social": [("B1", [0,2]), ("B2", [0,1,2]), ("B3", [0,1,2,7,8]), ("B4", [0,1,2,3]), ("B5", [0,1,2,3]), ("B6", [0,1,2]), ("B7", [0,1,2,3]), ("B8", [0,1,2]), ("B9", [0,1,2,3]), ("B10", [0,1,2,3,7]), ("B11", [0,1,2,3]), ("B12", [0,1,2,3]), ("B13", [0,1,2,3])],
    "C. Imaginación": [("C1", [0,1,2,3])],
    "D. Repetitivos": [("D1", [0,1,2,3]), ("D2", [0,1,2,3]), ("D3", [0,1,2,3]), ("D4", [0,1,2,3]), ("D5", [0,1,2])],
    "E. Otros": [("E1", [0,1,2,3,7]), ("E2", [0,1,2,3]), ("E3", [0,1,2])]
}

# --- TEXTOS CLÍNICOS ---
# Reescritos en registro de informe profesional: frases completas, con concordancia
# de género/sujeto mediante {art}/{suj}/{pron}, y un léxico clínico más preciso.
# El significado clínico de cada puntuación se mantiene igual que en el original;
# solo cambia la redacción.
TEXTOS = {
    # A. Lenguaje y comunicación
    "A1": {
        0: "{art} construye oraciones gramaticalmente correctas y de una complejidad adecuada.",
        1: "utiliza un habla relativamente compleja, aunque con errores gramaticales ocasionales.",
        2: "su habla se limita a expresiones de al menos tres palabras, sin llegar a una construcción compleja.",
        3: "su habla consiste principalmente en frases simples y de corta longitud.",
    },
    "A2": {
        0: "la entonación, el volumen y el ritmo del habla resultan adecuados al contexto.",
        1: "presenta una entonación algo monótona, o un volumen y ritmo levemente inusuales.",
        2: "el habla resulta claramente anómala, con una cualidad mecánica o un ritmo irregular.",
        7: "presenta alteraciones en la fluidez verbal, con episodios de tartamudeo.",
    },
    "A4": {
        0: "no emplea palabras ni frases estereotipadas.",
        1: "su lenguaje tiende a ser algo repetitivo o de un registro excesivamente formal.",
        2: "emplea vocalizaciones estereotipadas o expresiones idiosincrásicas.",
        3: "su habla es, casi en su totalidad, de carácter estereotipado.",
    },
    "A8": {
        0: "mantiene una conversación recíproca y fluida, construyendo el intercambio sobre las aportaciones del interlocutor.",
        1: "aporta elaboraciones espontáneas, si bien la cantidad global de habla resulta limitada.",
        2: "muestra escasa reciprocidad conversacional y tiende a seguir su propio hilo de pensamiento.",
        3: "presenta muy poca habla comunicativa de carácter espontáneo.",
    },
    "A9": {
        0: "emplea diversos gestos descriptivos, bien coordinados con el habla.",
        1: "utiliza gestos poco variados o recurre con frecuencia a gestos de tipo instrumental.",
        2: "hace un uso excepcional de gestos informativos, de forma aislada.",
        3: "presenta ausencia o un uso muy limitado de gestos.",
        8: "no procede valorar este aspecto.",
    },
    "A10": {
        0: "utiliza una variedad de gestos enfáticos o emocionales, bien integrados con el discurso.",
        1: "presenta gestos enfáticos limitados en frecuencia o en su integración con el habla.",
        2: "sus gestos enfáticos resultan llamativos o claramente mal integrados.",
        3: "presenta escasos o nulos gestos enfáticos.",
        8: "no procede valorar este aspecto.",
    },

    # B. Interacción social recíproca
    "B1": {
        0: "{art} mantiene un contacto visual socialmente modulado y flexible.",
        2: "establece un contacto visual pobremente modulado.",
    },
    "B2": {
        0: "dirige al examinador diversas expresiones faciales apropiadas.",
        1: "dirige algunas expresiones faciales, aunque con escasa variedad.",
        2: "no dirige expresiones faciales apropiadas al examinador.",
    },
    "B6": {
        0: "demuestra una clara comprensión e identificación de las emociones ajenas.",
        1: "transmite cierta comprensión e identifica alguna emoción ajena, de forma parcial.",
        2: "muestra escasa o nula identificación de las emociones de los demás.",
    },
    "B8": {
        0: "se describe a sí mism{term} como responsable de sus propias acciones.",
        1: "ofrece indicios claros de sentido de la responsabilidad, aunque de forma poco consistente.",
        2: "muestra escasos indicios de un sentido de responsabilidad sobre sus acciones.",
    },
    "B9": {
        0: "realiza iniciaciones sociales efectivas y claramente dirigidas al interlocutor.",
        1: "sus iniciaciones resultan inusuales o se centran predominantemente en sus propios intereses.",
        2: "sus iniciaciones son inapropiadas o carecen de una finalidad social clara.",
        3: "no realiza iniciaciones sociales.",
    },
    "B12": {
        0: "en conjunto, su comunicación social es extensa y recíproca.",
        1: "su comunicación recíproca se muestra reducida en frecuencia.",
        2: "su comunicación se orienta preferentemente hacia objetos o preocupaciones propias.",
        3: "presenta escasa o nula comunicación social.",
    },

    # C. Imaginación
    "C1": {
        0: "{art} muestra creatividad y aporta comentarios originales a lo largo de la evaluación.",
        1: "sus acciones creativas resultan poco variadas.",
        2: "su imaginación se muestra limitada o de carácter repetitivo.",
        3: "no se observan acciones de carácter creativo.",
    },

    # D. Comportamientos repetitivos e intereses restringidos
    "D1": {
        0: "no se observan intereses sensoriales inusuales.",
        1: "se aprecian algunos posibles intereses sensoriales, sin especial relevancia clínica.",
        2: "muestra un interés evidente por elementos sensoriales del entorno.",
        3: "presenta una búsqueda sensorial inusual y marcada.",
    },
    "D2": {
        0: "no presenta manierismos de manos o dedos.",
        1: "presenta manierismos poco claros o de dudosa significación.",
        2: "se observan retorcimientos de dedos de forma evidente.",
        3: "presenta manierismos frecuentes a lo largo de la evaluación.",
    },
    "D3": {
        0: "no muestra ningún intento de autolesión.",
        1: "se observan indicios dudosos de conducta autolesiva.",
        2: "se registran ejemplos claros de conducta autolesiva.",
        3: "presenta autolesiones de carácter grave.",
    },
    "D4": {
        0: "no muestra intereses excesivos hacia temas inusuales.",
        1: "realiza referencias ocasionales a intereses inusuales.",
        2: "presenta un patrón de intereses de carácter estereotipado.",
        3: "sus preocupaciones interfieren de forma significativa con su funcionamiento.",
    },
    "D5": {
        0: "no se observan compulsiones ni rituales.",
        1: "muestra actividades marcadamente fijadas a una rutina.",
        2: "presenta rituales con un componente compulsivo evidente.",
    },

    # E. Otros comportamientos
    "E1": {
        0: "mantiene un nivel de actividad apropiado.",
        1: "se muestra ligeramente inquiet{term}.",
        2: "presenta dificultades para permanecer sentad{term}.",
        3: "muestra un nivel de hiperactividad marcado.",
        7: "se muestra excesivamente quiet{term}, con la actividad notablemente reducida.",
    },
    "E2": {
        0: "no se muestra enfadad{term} ni presenta comportamientos disruptivos.",
        1: "presenta una ligera disrupción conductual.",
        2: "presenta más de un comportamiento disruptivo a lo largo de la sesión.",
        3: "presenta berrinches o conductas de agresión.",
    },
    "E3": {
        0: "no se observa ansiedad evidente.",
        1: "muestra signos leves de ansiedad.",
        2: "presenta una ansiedad marcada.",
    },
}

# 3. INTERFAZ Y RECOLECCIÓN DE DATOS
st.title("🔬 Sistema de Evaluación ADOS-2 Módulo 4")

with st.sidebar:
    gen = st.radio("Género:", ["Varón", "Mujer"])
    edad = st.number_input("Edad:", 12, 99, 25)
    if gen == "Varón":
        VARS = {"art": "El evaluado", "suj": "el evaluado", "pron": "él", "Pron": "Él", "term": "o"}
    else:
        VARS = {"art": "La evaluada", "suj": "la evaluada", "pron": "ella", "Pron": "Ella", "term": "a"}
    art = VARS["art"]
    del_al = "del evaluado" if gen == "Varón" else "de la evaluada"

tabs_names = ["Tareas"] + list(UI_STRUCTURE.keys())
tabs = st.tabs(tabs_names)

notas_tareas = {}
with tabs[0]:
    for t in TASKS:
        notas_tareas[t] = st.text_input(f"Notas breves: {t}", key=f"task_{t}")

scores = {}
for i, (section, items) in enumerate(UI_STRUCTURE.items()):
    with tabs[i + 1]:
        for key, options in items:
            label = f"{key}*" if key in ALGO_KEYS else key
            scores[key] = st.radio(label, options, horizontal=True)

# 4. GENERACIÓN DEL INFORME
if st.button("🚀 GENERAR INFORME TOTAL"):

    sc_C = sum(c(scores[k]) for k in ["A4", "A8", "A9", "A10"])
    sc_ISR = sum(c(scores[k]) for k in ["B1", "B2", "B6", "B8", "B9", "B11", "B12"])
    total = sc_C + sc_ISR

    if sc_C >= 3 and sc_ISR >= 6 and total >= 10:
        clima = "Autismo"
    elif sc_C >= 2 and sc_ISR >= 4 and total >= 7:
        clima = "Espectro del autismo"
    else:
        clima = "No TEA"

    def get_text(k):
        """Obtiene el texto base y sustituye las variables dinámicas de género/sujeto."""
        raw = TEXTOS.get(k, {}).get(scores[k], "")
        return raw.format(**VARS) if raw else raw

    def frase(conector, texto):
        """
        Combina un conector (en minúscula, puede ir vacío) con el texto clínico
        y devuelve una oración autónoma: primera letra en mayúscula y punto
        final. Cada llamada produce una oración completa, de modo que al
        unir varias con espacios el resultado son oraciones bien separadas
        (y no un único punto seguido de minúscula).
        """
        s = (conector + texto).strip()
        if not s:
            return ""
        if not s.endswith((".", "!", "?")):
            s += "."
        return s[0].upper() + s[1:]

    def parrafo(frases):
        """Une oraciones ya formadas (ver frase()) en un único párrafo."""
        return " ".join(f for f in frases if f)

    # --- Construcción narrativa de cada área, con conectores variados
    #     para evitar el efecto de "checklist" del listado de frases sueltas.
    #     Cada línea es una oración completa por sí misma. ---

    redA = parrafo([
        frase("", get_text("A1")),
        frase("Por lo que respecta a la prosodia, ", get_text("A2")),
        frase("En el plano del uso del lenguaje, ", get_text("A4")),
        frase("En el terreno conversacional, ", get_text("A8")),
        frase("A nivel gestual, ", get_text("A9")),
        frase("En cuanto a los gestos enfáticos, ", get_text("A10")),
    ])

    redB = parrafo([
        frase("", get_text("B1")),
        frase("", get_text("B2")),
        frase("Por otro lado, ", get_text("B6")),
        frase("En cuanto a la conducta social, ", get_text("B8")),
        frase("", get_text("B9")),
        frase("", get_text("B12")),
    ])

    redC = parrafo([
        frase("", get_text("C1")),
    ])

    redD = parrafo([
        frase("", get_text("D1")),
        frase("Asimismo, ", get_text("D2")),
        frase("Por otro lado, ", get_text("D3")),
        frase("En cuanto a sus intereses, ", get_text("D4")),
        frase("Finalmente, en lo relativo a rituales y rutinas, ", get_text("D5")),
    ])

    redE = parrafo([
        frase("Durante la evaluación, ", get_text("E1")),
        frase("", get_text("E2")),
        frase("En cuanto al estado emocional, ", get_text("E3")),
    ])

    intro_doc = (
        f"El ADOS-2 (Autism Diagnostic Observation Schedule, segunda edición) es una escala de "
        f"observación semiestructurada, estandarizada y validada internacionalmente, diseñada para "
        f"la evaluación de comportamientos asociados al espectro autista. El Módulo 4, aplicado en "
        f"esta evaluación, está indicado para adolescentes mayores y personas adultas con lenguaje fluido.\n\n"
        f"A lo largo de la aplicación se plantearon diversas situaciones estructuradas y "
        f"semiestructuradas que permitieron observar la comunicación, la interacción social recíproca, "
        f"la imaginación/creatividad y la presencia de intereses restringidos o comportamientos "
        f"repetitivos {del_al}. A continuación se describe el perfil observado en cada una de las "
        f"áreas evaluadas."
    )

    resultados_texto = (
        f"Los resultados obtenidos en el algoritmo diagnóstico del Módulo 4 son los siguientes: "
        f"puntuación de Comunicación = {sc_C}; puntuación de Interacción social recíproca = {sc_ISR}; "
        f"puntuación total = {total}. Clasificación clínica resultante: {clima}."
    )

    informe_final = (
        f"APLICACIÓN DEL ADOS-2\n\n{intro_doc}\n\n"
        f"1. COMUNICACIÓN\n{redA}\n\n"
        f"2. INTERACCIÓN SOCIAL\n{redB}\n\n"
        f"3. IMAGINACIÓN\n{redC}\n\n"
        f"4. COMPORTAMIENTOS REPETITIVOS E INTERESES RESTRINGIDOS\n{redD}\n\n"
        f"5. OTROS COMPORTAMIENTOS\n{redE}\n\n"
        f"RESULTADOS DEL ALGORITMO\n{resultados_texto}"
    )

    st.text_area("Vista previa del Informe:", informe_final, height=500)

    # --- Generación del documento Word con estructura real (títulos y
    #     párrafos separados), en lugar de un único bloque de texto ---
    doc = Document()
    doc.add_heading("INFORME ADOS-2 (MÓDULO 4)", 0)

    doc.add_heading("Aplicación del ADOS-2", level=1)
    for parte in intro_doc.split("\n\n"):
        doc.add_paragraph(parte)

    secciones = [
        ("1. Comunicación", redA),
        ("2. Interacción social", redB),
        ("3. Imaginación", redC),
        ("4. Comportamientos repetitivos e intereses restringidos", redD),
        ("5. Otros comportamientos", redE),
    ]
    for titulo, texto in secciones:
        doc.add_heading(titulo, level=1)
        doc.add_paragraph(texto)

    doc.add_heading("Resultados del algoritmo", level=1)
    doc.add_paragraph(f"Comunicación: {sc_C}", style="List Bullet")
    doc.add_paragraph(f"Interacción social recíproca: {sc_ISR}", style="List Bullet")
    doc.add_paragraph(f"Puntuación total: {total}", style="List Bullet")
    p = doc.add_paragraph()
    p.add_run(f"Clasificación clínica: {clima}").bold = True

    notas_activas = {k: v for k, v in notas_tareas.items() if v.strip()}
    if notas_activas:
        doc.add_heading("Observaciones específicas por tarea", level=1)
        for task, note in notas_activas.items():
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{task}: ").bold = True
            para.add_run(note)

    bio = BytesIO()
    doc.save(bio)
    st.download_button(
        "📥 Descargar Documento Word",
        bio.getvalue(),
        f"Informe_ADOS2_{gen}_{edad}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
